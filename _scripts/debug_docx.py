"""Debug script to inspect Word document structure."""

import sys
from pathlib import Path
from docx import Document

def debug_document(docx_path: Path):
    """Print document structure for debugging."""
    print(f"Opening: {docx_path}\n")
    doc = Document(str(docx_path))

    print("="*60)
    print("PARAGRAPHS (first 100)")
    print("="*60)
    for i, para in enumerate(doc.paragraphs[:100]):
        text = para.text.strip()
        if text:
            print(f"{i:3d}: {text[:80]}")
            if '7.' in text and text.startswith('7.'):
                print(f"      ^^^ POTENTIAL SECTION 7 MATCH ^^^")

    print(f"\n{'='*60}")
    print(f"TABLES: {len(doc.tables)} found")
    print("="*60)
    for i, table in enumerate(doc.tables[:10]):  # Show first 10 tables
        print(f"\nTable {i}:")
        print(f"  Rows: {len(table.rows)}")
        print(f"  Columns: {len(table.columns) if table.rows else 0}")
        if table.rows:
            print(f"  First row (headers):")
            for cell in table.rows[0].cells[:5]:  # First 5 cells
                print(f"    - {cell.text.strip()[:40]}")

if __name__ == '__main__':
    if len(sys.argv) < 2:
        print("Usage: python debug_docx.py <path_to_docx>")
        sys.exit(1)

    debug_document(Path(sys.argv[1]))
