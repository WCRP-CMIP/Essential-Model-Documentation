"""Extract CV terms from EMD Word document examples."""

import argparse
import re
import sys
from pathlib import Path
from typing import Dict, List, Optional, Set

from docx import Document

from utils import (
    normalize_id,
    generate_drs_name,
    write_json,
    get_descriptor_from_section
)


def extract_cv_reference(text: str) -> Optional[tuple[str, str]]:
    """
    Extract CV reference from text.

    Args:
        text: Paragraph text

    Returns:
        Tuple of (section_number, cv_name) or None
        E.g., ("7.1", "component") from "7.1 (component CV)"
    """
    # Match patterns like "7.1 (component CV)" or "7.2 (calendar CV)"
    match = re.search(r'(7\.\d+)\s*\((\w+(?:_\w+)?)\s+CV\)', text)
    if match:
        return (match.group(1), match.group(2))
    return None


def parse_example_terms(text: str) -> List[str]:
    """
    Parse terms from example paragraph.

    Args:
        text: Paragraph text starting with "E.g."

    Returns:
        List of term strings
    """
    # Remove "E.g." prefix
    text = re.sub(r'^E\.g\.\s*', '', text, flags=re.IGNORECASE).strip()

    # Split by comma or "and"
    terms = re.split(r',|\band\b', text)

    # Clean up each term
    cleaned_terms = []
    for term in terms:
        term = term.strip()
        # Remove any trailing punctuation
        term = re.sub(r'[.,;:]$', '', term)
        if term:
            cleaned_terms.append(term)

    return cleaned_terms


def extract_cvs_from_docx(docx_path: Path) -> Dict[str, Dict]:
    """
    Extract all CV terms from EMD Word document examples.

    Args:
        docx_path: Path to Word document

    Returns:
        Dictionary mapping descriptor names to CV data
    """
    print(f"Opening Word document: {docx_path}")
    doc = Document(str(docx_path))

    # Dictionary to collect CV references and their terms
    # Format: {section_num: {'cv_name': str, 'terms': Set[str]}}
    cv_data = {}

    # First pass: find all CV references and collect example terms
    print("\nSearching for CV references and example terms...")

    # Track CV references with their paragraph indices
    cv_references = {}  # {paragraph_index: (section_num, cv_name)}

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()

        # Check for CV reference (e.g., "7.1 (component CV)")
        cv_ref = extract_cv_reference(text)
        if cv_ref:
            section_num, cv_name = cv_ref
            cv_references[i] = cv_ref

            if section_num not in cv_data:
                cv_data[section_num] = {
                    'cv_name': cv_name,
                    'terms': set(),
                    'title': cv_name.replace('_', ' ').title()
                }
                print(f"  Found CV: {section_num} ({cv_name} CV)")

    # Second pass: collect example terms that are close to CV references
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()

        # Check for example terms (E.g. ...)
        if text.startswith('E.g.') or text.startswith('e.g.'):
            # Look for the most recent CV reference within last 5 paragraphs
            found_cv = None
            for j in range(max(0, i - 5), i):
                if j in cv_references:
                    found_cv = cv_references[j]
                    break

            if found_cv:
                section_num, _ = found_cv
                # Parse and add terms
                terms = parse_example_terms(text)
                # Filter out non-CV terms (numbers, years, very long phrases)
                valid_terms = []
                for term in terms:
                    # Skip if it's just a number or year
                    if term.isdigit() or (len(term) == 4 and term.isdigit()):
                        continue
                    # Skip very long phrases (likely not CV terms)
                    if len(term) > 50:
                        continue
                    valid_terms.append(term)

                if valid_terms:
                    for term in valid_terms:
                        cv_data[section_num]['terms'].add(term)
                    print(f"    Added {len(valid_terms)} example terms to {section_num}")

    # Convert to final format
    cvs = {}
    for section_num, data in sorted(cv_data.items()):
        if not data['terms']:
            print(f"  Warning: No terms found for {section_num} ({data['cv_name']})")
            continue

        # Get descriptor name from mapping
        descriptor_name = get_descriptor_from_section(section_num)
        if not descriptor_name:
            descriptor_name = data['cv_name']
            print(f"  Warning: No mapping for {section_num}, using cv_name: {descriptor_name}")

        # Create term entries
        terms = []
        for term_name in sorted(data['terms']):
            term_id = normalize_id(term_name)
            terms.append({
                'id': term_id,
                'description': f'{term_name}',  # Use original name as description
                'drs_name': generate_drs_name(term_id)
            })

        cvs[descriptor_name] = {
            'section': section_num,
            'title': data['title'],
            'descriptor_name': descriptor_name,
            'terms': terms
        }

        print(f"\n{section_num} ({descriptor_name}): {len(terms)} terms")
        for term in terms[:5]:  # Show first 5 terms
            print(f"    - {term['id']}")
        if len(terms) > 5:
            print(f"    ... and {len(terms) - 5} more")

    return cvs


def main():
    """Main entry point."""
    parser = argparse.ArgumentParser(description='Extract CV tables from EMD Word document')
    parser.add_argument('--docx', required=True, help='Path to EMD Word document')
    parser.add_argument('--output', required=True, help='Output JSON file path')

    args = parser.parse_args()

    docx_path = Path(args.docx)
    output_path = Path(args.output)

    if not docx_path.exists():
        print(f"Error: Word document not found: {docx_path}")
        sys.exit(1)

    # Extract CVs
    print("\nExtracting CVs from Word document...")
    cvs = extract_cvs_from_docx(docx_path)

    if not cvs:
        print("Warning: No CVs extracted from document")
        sys.exit(1)

    # Write output
    print(f"\nWriting extracted CVs to: {output_path}")
    write_json(output_path, cvs)

    # Print summary
    print("\n=== Extraction Summary ===")
    print(f"Total descriptors extracted: {len(cvs)}")
    for descriptor_name, cv_data in sorted(cvs.items()):
        print(f"  {descriptor_name}: {len(cv_data['terms'])} terms (Section {cv_data['section']})")

    print("\nExtraction complete!")


if __name__ == '__main__':
    main()
